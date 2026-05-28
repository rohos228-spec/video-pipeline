#!/usr/bin/env python3
"""Generate Word doc: plan for Excel batch video pipeline."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("План: таблица Excel → много видео\nс общими настройками промптов")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"video-pipeline · {date.today().strftime('%d.%m.%Y')}")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    add_para(
        doc,
        "Документ описывает логику массового производства видео: вместо одной темы "
        "загружается таблица, где каждая строка — отдельное видео. "
        "Пайплайн сохраняет настройки промптов, но для каждого видео начинает работу заново "
        "в отдельной папке.",
    )

    add_heading(doc, "1. Цель", 1)
    add_para(doc, "Настроить пайплайн один раз (промпты, стили, генераторы, canvas). "
             "Загрузить Excel, где каждая строка = отдельное видео. Для каждой строки система:")
    add_bullets(doc, [
        "берёт тему и поля из строки;",
        "сохраняет выбранные промпты и переменные;",
        "сбрасывает прогресс пайплайна (как «новое видео с нуля»);",
        "создаёт отдельную папку data/videos/<slug>/.",
    ])

    add_heading(doc, "2. Что уже есть в коде (~70%)", 1)
    add_para(doc, "В Studio уже реализована «Массовая генерация» (excel_feed + mass-lanes/start):")
    add_table(doc,
        ["Уже работает", "Чего не хватает"],
        [
            ["Загрузка topics.xlsx на canvas", "В мастере «Новый проект» только поле «тема»"],
            ["Каждая строка → отдельный Project + папка", "Из Excel на web берётся только название (кол. B)"],
            ["Копируются prompt_overrides, gpt_text_overrides", "meta.prompt_styles не подключены к бэкенду"],
            ["Очередь: одно видео за раз", "Нет явного UI «что сохраняется / что сбрасывается»"],
            ["Telegram-батч с полной карточкой строки", "Web и Telegram — два разных механизма"],
        ],
    )
    add_para(doc, "Вывод: не писать с нуля, а доработать mass-lanes и добавить понятный вход через таблицу.", bold=True)

    add_heading(doc, "3. Три слоя данных", 1)

    add_heading(doc, "3.1. Сохраняется навсегда (шаблон / родительский проект)", 2)
    add_bullets(doc, [
        "Выбранные варианты промптов (prompt_overrides, prompt_slot_variants)",
        "Переменные промптов (vars, blocks, style_profile в prompt_overrides)",
        "GPT-тексты нод (gpt_text_overrides)",
        "Настройки: генераторы, aspect, relax, hero_count, enrich_slots",
        "Canvas: disabled_nodes, custom_prompts",
        "Сам файл таблицы и привязки excel → plan",
    ])

    add_heading(doc, "3.2. Сбрасывается на каждое новое видео", 2)
    add_bullets(doc, [
        "Статус пайплайна → new",
        "Кадры (Frames), артефакты (Artifacts), HITL-запросы",
        "Пустой project.xlsx только для этого видео",
        "Пустые подпапки: characters/, scenes/, videos/, audio/, final/",
    ])

    add_heading(doc, "3.3. Берётся из строки Excel", 2)
    add_bullets(doc, [
        "Колонка B — название / тема",
        "Колонки C–J — карточка (стиль, хук, факт…) → meta.topic_card и промпт плана",
        "Колонка K — hero_mode для этой строки (если задан)",
    ])

    add_heading(doc, "4. Пользовательский сценарий (UX)", 1)

    add_heading(doc, "Шаг 1 — Создать «фабрику» (шаблон)", 2)
    add_bullets(doc, [
        "Новый проект → режим «Из таблицы» (вместо одной темы).",
        "Загрузить topics.xlsx или создать пустой шаблон.",
        "Настроить canvas и промпты на родительском проекте.",
    ])

    add_heading(doc, "Шаг 2 — Настроить промпты один раз", 2)
    add_bullets(doc, [
        "В каждой ноде выбрать файлы промптов, стили, переменные.",
        "Настройки живут на родителе и копируются в каждое дочернее видео при старте очереди.",
    ])

    add_heading(doc, "Шаг 3 — Запустить очередь", 2)
    add_bullets(doc, [
        "Кнопка «Запустить N видео из таблицы».",
        "Для каждой строки: дочерний проект + папка, копия настроек промптов, чистый pipeline, старт «План».",
    ])

    add_heading(doc, "Шаг 4 — Мониторинг", 2)
    add_bullets(doc, [
        "Родитель: таблица со статусами строк (slug, статус, прогресс).",
        "Дочерние проекты в сайдбаре (фильтр «из очереди X»).",
        "Очередь последовательная — одно видео в работе (как сейчас).",
    ])

    add_heading(doc, "5. Формат таблицы", 1)
    add_para(doc, "Используется topics.xlsx (лист «Темы»), как в Telegram-батче. "
             "Не project.xlsx (там колонки = кадры одного ролика).")
    add_table(doc,
        ["Колонка", "Смысл", "Обязательно"],
        [
            ["B", "Название ролика", "Да"],
            ["C–J", "Карточка (стиль, хук, факт, интеграция продукта…)", "Нет"],
            ["K", "hero_mode (hero / no_hero / auto)", "Нет"],
            ["L–O", "slug, статус, прогресс, дата обновления", "Пишет система"],
        ],
    )

    add_heading(doc, "6. Технический план по этапам", 1)

    add_heading(doc, "Этап A — Бэкенд: полная строка Excel → видео", 2)
    add_bullets(doc, [
        "Файлы: app/web/routers/project_ops.py, app/storage/batch_sheet.py",
        "При mass-lanes/start передавать cards[] (не только topics: string[]).",
        "В дочерний проект писать meta.topic_card из строки.",
        "Подключить topic_card в промпт плана (как в TG batches.py).",
        "После завершения видео обновлять строку в Excel родителя (статус, slug).",
    ])

    add_heading(doc, "Этап B — Что копируем / не копируем", 2)
    add_bullets(doc, [
        "Новый модуль app/services/mass_lane_clone.py с явным whitelist/blacklist.",
        "COPY: prompt_overrides, gpt_text_overrides, generators, hero_*, meta.prompt_*",
        "STRIP: mass_queue_*, mass_excel_*, excel_lane_bindings, frames, artifacts",
        "FROM_ROW: topic, topic_card, hero_mode (override)",
    ])

    add_heading(doc, "Этап C — Промпт-стили реально работают", 2)
    add_bullets(doc, [
        "Файлы: app/services/prompt_composer.py, web/node-studio.tsx",
        "Объединить meta.prompt_styles → prompt_overrides при сохранении.",
        "Иначе «стили промптов сохраняются» только в UI, но не в генерации.",
    ])

    add_heading(doc, "Этап D — UI", 2)
    add_bullets(doc, [
        "Мастер: переключатель «Одна тема» / «Таблица видео».",
        "Excel feed: показ всех колонок карточки.",
        "Панель очереди на родителе: N строк, статус, старт/пауза.",
        "Подсказка: «Настройки промптов → применятся ко всем строкам».",
    ])

    add_heading(doc, "Этап E — Опционально: постоянные ассеты", 2)
    add_bullets(doc, [
        "Если нужен один продукт/герой во всех видео (как permanent_product в TG).",
        "Поле в родителе: прикрепить референс.",
        "При fork копировать только этот артефакт.",
    ])

    add_heading(doc, "Этап F — Тесты", 2)
    add_bullets(doc, [
        "tests/test_mass_lanes_excel.py — карточки, topic_card, clone whitelist.",
        "Тест: fork не копируet frames/artifacts родителя.",
    ])

    add_heading(doc, "7. Варианты архитектуры", 1)
    add_table(doc,
        ["Вариант", "Суть", "Плюсы", "Минусы", "Рекомендация"],
        [
            ["A. Доработать mass-lanes", "Родитель-шаблон + N детей", "Мало риска, уже работает", "Много проектов в sidebar", "Да, v1"],
            ["B. Один проект + смена строки", "current_row, reset", "Проще в списке", "Сложный reset", "Нет"],
            ["C. Объединить с TG BatchProject", "Одна модель БД", "Один код", "Большой рефакторинг", "Позже, v2"],
        ],
    )

    add_heading(doc, "8. Что не делаем в первой версии", 1)
    add_bullets(doc, [
        "Разные subgraph canvas на каждую строку.",
        "Параллельная генерация нескольких видео сразу.",
        "Импорт project.xlsx как списка видео.",
    ])

    add_heading(doc, "9. Вопросы для уточнения перед реализацией", 1)
    add_para(doc, "«Сохранять переменные или результаты генерации» — два режима:")
    add_table(doc,
        ["Режим", "Смысл", "Пример"],
        [
            ["A — только настройки", "Копируем промпты/vars, не картинки/видео", "30 роликов, разные темы, один стиль промптов"],
            ["B — + общие ассеты", "Плюс один продукт/герой/референс", "Все ролики с одним продуктом в кадре"],
        ],
    )
    add_bullets(doc, [
        "Режим A заложен в v1; режим B — опционально на этапе E.",
        "Таблица — всегда topics.xlsx (строка = видео) или любой xlsx?",
    ])

    add_heading(doc, "10. Ключевые файлы кодовой базы", 1)
    add_table(doc,
        ["Область", "Путь"],
        [
            ["Модели БД", "app/models.py"],
            ["Mass lanes API", "app/web/routers/project_ops.py"],
            ["TG batch", "app/services/batches.py"],
            ["Парсинг topics.xlsx", "app/storage/batch_sheet.py"],
            ["Очередь", "app/orchestrator/auto_advance.py"],
            ["Промпты", "app/services/prompt_composer.py, prompt_library.py"],
            ["Canvas / Excel feed", "web/src/components/canvas/flow-canvas.tsx, excel-feed-panel.tsx"],
            ["Мастер проекта", "web/src/components/sidebar/new-project-wizard.tsx"],
            ["Документация mass", "docs/MASS_CREATION.md"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Статус: план / согласование · реализация не начата")
    fr.italic = True
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def main() -> None:
    out_paths = [
        Path("/workspace/docs/plans/2026-05-28-excel-batch-video-plan.docx"),
        Path("/opt/cursor/artifacts/План-Excel-массовые-видео.docx"),
    ]
    doc = build_doc()
    for p in out_paths:
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        print(f"Saved: {p} ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
